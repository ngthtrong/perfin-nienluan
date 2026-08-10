#!/usr/bin/env python3
"""Export the English PERFIN LaTeX report to a guideline-compliant DOCX.

The repository intentionally keeps LaTeX as the source of truth.  This exporter
handles the small, known LaTeX vocabulary used by the report and creates native
Word headings, tables, captions, TOC/list fields, page numbering, and images.
Display equations are rendered from the original TeX, so their mathematical
content is not approximated by a text-only conversion.
"""

from __future__ import annotations

import argparse
import re
import struct
import subprocess
import tempfile
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
EN_FILES = [
    ROOT / "frontmatter/en/abstract.tex",
    ROOT / "frontmatter/en/abbreviations.tex",
    ROOT / "chapters/en/chapter1.tex",
    ROOT / "chapters/en/chapter2.tex",
    ROOT / "chapters/en/chapter3.tex",
    ROOT / "chapters/en/chapter4.tex",
    ROOT / "chapters/en/references.tex",
    ROOT / "chapters/en/appendices.tex",
]

META = {
    "university": "CAN THO UNIVERSITY",
    "college": "COLLEGE OF INFORMATION AND COMMUNICATION TECHNOLOGY",
    "projecttype": "Project – Fundamental Topics",
    "projecttitle": (
        "A MOBILE PERSONAL FINANCE MANAGEMENT APPLICATION SUPPORTED BY "
        "A LARGE LANGUAGE MODEL — PERFIN"
    ),
    "major": "Software Engineering",
    "cohort": "49",
    "courseclass": "CT239H M01",
    "advisor": "Dr. Phan Phuong Lan",
    "studentname": "Nguyen Thanh Trong",
    "studentid": "B2305615",
    "semester": "3",
    "academicyear": "2025–2026",
}

HEADING_ONE = "00B0F0"
HEADING_TWO = "2F5496"


def strip_comments(text: str) -> str:
    """Remove unescaped TeX comments without touching escaped percent signs."""
    out = []
    for line in text.splitlines():
        escaped = False
        keep = []
        for ch in line:
            if ch == "%" and not escaped:
                break
            keep.append(ch)
            if ch == "\\" and not escaped:
                escaped = True
            else:
                escaped = False
        out.append("".join(keep).rstrip())
    return "\n".join(out)


def brace_arg(text: str, pos: int) -> tuple[str, int]:
    """Read one balanced {...} argument starting at or after *pos*."""
    while pos < len(text) and text[pos].isspace():
        pos += 1
    if pos >= len(text) or text[pos] != "{":
        return "", pos
    depth = 1
    start = pos + 1
    pos += 1
    escaped = False
    while pos < len(text):
        ch = text[pos]
        if ch == "{" and not escaped:
            depth += 1
        elif ch == "}" and not escaped:
            depth -= 1
            if depth == 0:
                return text[start:pos], pos + 1
        escaped = ch == "\\" and not escaped
        if ch != "\\":
            escaped = False
        pos += 1
    raise ValueError("Unbalanced TeX argument")


def command_args(line: str, command: str, count: int) -> list[str]:
    pos = line.find("\\" + command)
    if pos < 0:
        return []
    pos += len(command) + 1
    args = []
    for _ in range(count):
        arg, pos = brace_arg(line, pos)
        args.append(arg)
    return args


@dataclass
class ReferenceMap:
    labels: dict[str, str] = field(default_factory=dict)
    citations: dict[str, int] = field(default_factory=dict)
    table_captions: list[str] = field(default_factory=list)
    figure_captions: list[str] = field(default_factory=list)


def build_reference_map() -> ReferenceMap:
    """Pre-scan the source so forward references can be emitted as plain text."""
    result = ReferenceMap()
    chapter = section = subsection = subsubsection = 0
    appendix = False
    appendix_index = 0
    table_no = figure_no = 0
    pending_ref: str | None = None

    token_re = re.compile(
        r"\\appendix\b|"
        r"\\(?:chapter|section|subsection|subsubsection)\{|"
        r"\\begin\{frtable\}|"
        r"\\caption\{|"
        r"\\widereportfigure\{|"
        r"\\label\{|"
        r"\\bibitem\{"
    )

    for path in EN_FILES:
        text = strip_comments(path.read_text(encoding="utf-8"))
        pos = 0
        while True:
            match = token_re.search(text, pos)
            if not match:
                break
            token = match.group(0)
            start = match.start()
            pos = match.end()
            if token.startswith("\\appendix"):
                appendix = True
                pending_ref = None
                continue
            if token.startswith("\\bibitem"):
                key, pos = brace_arg(text, pos - 1)
                result.citations[key] = len(result.citations) + 1
                continue
            if token.startswith("\\chapter"):
                _, pos = brace_arg(text, pos - 1)
                section = subsection = subsubsection = 0
                if appendix:
                    appendix_index += 1
                    pending_ref = chr(64 + appendix_index)
                else:
                    chapter += 1
                    pending_ref = str(chapter)
                continue
            if token.startswith("\\section"):
                _, pos = brace_arg(text, pos - 1)
                section += 1
                subsection = subsubsection = 0
                prefix = chr(64 + appendix_index) if appendix else str(chapter)
                pending_ref = f"{prefix}.{section}"
                continue
            if token.startswith("\\subsection"):
                _, pos = brace_arg(text, pos - 1)
                subsection += 1
                subsubsection = 0
                prefix = chr(64 + appendix_index) if appendix else str(chapter)
                pending_ref = f"{prefix}.{section}.{subsection}"
                continue
            if token.startswith("\\subsubsection"):
                _, pos = brace_arg(text, pos - 1)
                subsubsection += 1
                prefix = chr(64 + appendix_index) if appendix else str(chapter)
                pending_ref = f"{prefix}.{section}.{subsection}.{subsubsection}"
                continue
            if token.startswith("\\begin{frtable}"):
                caption, p1 = brace_arg(text, pos)
                label, pos = brace_arg(text, p1)
                table_no += 1
                result.table_captions.append(caption)
                result.labels[label] = str(table_no)
                pending_ref = str(table_no)
                continue
            if token.startswith("\\caption"):
                caption, pos = brace_arg(text, pos - 1)
                table_no += 1
                result.table_captions.append(caption)
                pending_ref = str(table_no)
                continue
            if token.startswith("\\widereportfigure"):
                filename, p1 = brace_arg(text, pos - 1)
                caption, p2 = brace_arg(text, p1)
                label, pos = brace_arg(text, p2)
                figure_no += 1
                result.figure_captions.append(caption)
                result.labels[label] = str(figure_no)
                pending_ref = str(figure_no)
                continue
            if token.startswith("\\label"):
                label, pos = brace_arg(text, pos - 1)
                if pending_ref:
                    result.labels[label] = pending_ref
                continue
    return result


def set_run_font(run, name: str, size: float | None = None) -> None:
    run.font.name = name
    if size is not None:
        run.font.size = Pt(size)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attr}"), name)


def set_style_font(style, name: str, size: float, color: str | None = None) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attr}"), name)


def add_outline_level(style, level: int) -> None:
    ppr = style.element.get_or_add_pPr()
    outline = ppr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        ppr.append(outline)
    outline.set(qn("w:val"), str(level))


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, "Times New Roman", 13)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(1)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.widow_control = True

    specs = [
        ("Heading 1", 14, True, False, HEADING_ONE, 0, 18, 8),
        ("Heading 2", 13, True, False, HEADING_TWO, 1, 14, 6),
        ("Heading 3", 13, True, True, HEADING_TWO, 2, 12, 5),
        ("Heading 4", 13, False, True, HEADING_TWO, 3, 10, 4),
    ]
    for name, size, bold, italic, color, level, before, after in specs:
        st = styles[name]
        set_style_font(st, "Arial", size, color)
        st.font.bold = bold
        st.font.italic = italic
        st.paragraph_format.first_line_indent = Cm(0)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
        add_outline_level(st, level)

    for name, listed in (("Front Heading 1", True), ("Unlisted Front Title", False)):
        if name not in styles:
            st = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            st = styles[name]
        st.base_style = normal
        set_style_font(st, "Arial", 14, HEADING_ONE)
        st.font.bold = True
        st.font.all_caps = True
        st.paragraph_format.first_line_indent = Cm(0)
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        st.paragraph_format.space_before = Pt(0)
        st.paragraph_format.space_after = Pt(18)
        st.paragraph_format.keep_with_next = True
        add_outline_level(st, 0 if listed else 9)

    caption = styles["Caption"]
    set_style_font(caption, "Times New Roman", 13)
    caption.font.italic = False
    caption.paragraph_format.first_line_indent = Cm(0)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(6)
    caption.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        st = styles[name]
        set_style_font(st, "Times New Roman", 13)
        st.paragraph_format.first_line_indent = Cm(0)
        st.paragraph_format.left_indent = Cm(1.1)
        st.paragraph_format.line_spacing = 1.5
        st.paragraph_format.space_after = Pt(2)

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    set_style_font(code, "DejaVu Sans Mono", 10.5)
    code.paragraph_format.first_line_indent = Cm(0)
    code.paragraph_format.left_indent = Cm(0.7)
    code.paragraph_format.right_indent = Cm(0.4)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(4)
    code.paragraph_format.line_spacing = 1.0


def set_page_geometry(section, landscape: bool = False) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    if landscape:
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
    else:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(3.5)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)


def set_page_number_format(section, fmt: str, start: int | None = None) -> None:
    sect_pr = section._sectPr
    node = sect_pr.find(qn("w:pgNumType"))
    if node is None:
        node = OxmlElement("w:pgNumType")
        sect_pr.append(node)
    node.set(qn("w:fmt"), fmt)
    if start is not None:
        node.set(qn("w:start"), str(start))
    elif qn("w:start") in node.attrib:
        del node.attrib[qn("w:start")]


def add_field(paragraph, instruction: str, result: str = "") -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    def append_field_run(node):
        run = OxmlElement("w:r")
        run.append(node)
        paragraph._p.append(run)

    append_field_run(begin)
    append_field_run(instr)
    append_field_run(separate)
    if result:
        run = paragraph.add_run(result)
        set_run_font(run, "Times New Roman", 13)
    append_field_run(end)


def add_page_number(section) -> None:
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    add_field(p, " PAGE ", "1")


def add_update_fields_setting(doc: Document) -> None:
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def add_cover_border(paragraph) -> None:
    """Add two independent VML page rectangles (bold outer, thin inner)."""
    xml = f"""
    <w:r {nsdecls('w')} xmlns:v="urn:schemas-microsoft-com:vml">
      <w:pict>
        <v:rect style="position:absolute;left:34pt;top:34pt;width:527pt;height:774pt;
          mso-position-horizontal-relative:page;mso-position-vertical-relative:page;
          z-index:-251654144" filled="f" strokecolor="#000000" strokeweight="2.25pt"/>
        <v:rect style="position:absolute;left:43pt;top:43pt;width:509pt;height:756pt;
          mso-position-horizontal-relative:page;mso-position-vertical-relative:page;
          z-index:-251654143" filled="f" strokecolor="#000000" strokeweight="0.75pt"/>
      </w:pict>
    </w:r>
    """
    paragraph._p.append(parse_xml(xml))


def add_cover(doc: Document) -> None:
    p = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    add_cover_border(p)
    for text in (META["university"], META["college"]):
        run = p.add_run(text + "\n")
        set_run_font(run, "Times New Roman", 14)
        run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(12)
    logo = ROOT / "images/ctu_logo.png"
    if logo.exists():
        p.add_run().add_picture(str(logo), width=Cm(3.3))
    else:
        run = p.add_run("[CTU LOGO]")
        run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run(META["projecttype"])
    set_run_font(run, "Times New Roman", 14)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(45)
    p.paragraph_format.space_after = Pt(38)
    run = p.add_run(META["projecttitle"])
    set_run_font(run, "Times New Roman", 18)
    run.bold = True

    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    labels = ["Major:", "Cohort:", "Course class:", "Advisor:", "Student:", "Student ID:"]
    values = [META["major"], META["cohort"], META["courseclass"], META["advisor"], META["studentname"], META["studentid"]]
    for row, label, value in zip(table.rows, labels, values):
        row.cells[0].width = Cm(5.2)
        row.cells[1].width = Cm(8.2)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].paragraph_format.first_line_indent = Cm(0)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        r = row.cells[0].paragraphs[0].add_run(label)
        set_run_font(r, "Times New Roman", 13)
        r.bold = True
        r = row.cells[1].paragraphs[0].add_run(value)
        set_run_font(r, "Times New Roman", 13)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(40)
    run = p.add_run(f"Semester {META['semester']}, academic year {META['academicyear']}")
    set_run_font(run, "Times New Roman", 13)
    run.bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run("Can Tho, 2026")
    set_run_font(run, "Times New Roman", 13)


MATH_COMMANDS = {
    "leq": "≤", "geq": "≥", "approx": "≈", "rightarrow": "→",
    "in": "∈", "subseteq": "⊆", "cap": "∩", "times": "×",
    "delta": "δ", "tau": "τ", "sigma": "σ", "sum": "Σ",
    "ldots": "…", "bar": "¯", "widehat": "^", "hat": "^",
    "lfloor": "⌊", "rfloor": "⌋", "lceil": "⌈", "rceil": "⌉",
    "quad": "  ", "qquad": "    ", "min": "min", "max": "max",
}


def math_to_text(value: str) -> str:
    text = " ".join(value.replace("\n", " ").split())
    text = re.sub(r"\\(?:left|right|displaystyle|substack)\b", "", text)
    text = re.sub(r"\\(?:,|!|;|:)\s*", " ", text)
    text = re.sub(r"\\bar\s*\{?([A-Za-z])\}?", lambda m: m.group(1) + "\u0304", text)
    text = re.sub(r"\\(?:widehat|hat)\s*\{?([A-Za-z])\}?", lambda m: m.group(1) + "\u0302", text)
    previous = None
    while previous != text:
        previous = text
        text = re.sub(
            r"\\(?:d?frac)\{([^{}]*)\}\{([^{}]*)\}",
            lambda m: f"({m.group(1)})/({m.group(2)})",
            text,
        )
        text = re.sub(r"\\sqrt\{([^{}]*)\}", r"√(\1)", text)
        text = re.sub(r"\\(?:text|mathrm|operatorname)\{([^{}]*)\}", r"\1", text)
    for command, replacement in MATH_COMMANDS.items():
        text = re.sub(rf"\\{command}(?![A-Za-z])", replacement, text)
    text = text.replace("\\_", "_").replace("\\%", "%")
    text = text.replace("\\{", "{").replace("\\}", "}")
    text = text.replace("\\", " ").replace("{", "").replace("}", "")
    return re.sub(r"\s+", " ", text).strip()


class InlineWriter:
    def __init__(self, refs: ReferenceMap):
        self.refs = refs

    def add(self, paragraph, text: str, *, bold=False, italic=False, code=False, color=None):
        text = text.strip()
        if not text:
            return
        self._walk(paragraph, text, bold=bold, italic=italic, code=code, color=color)

    def _plain(self, paragraph, text, *, bold, italic, code, color, math=False):
        if not text:
            return
        text = text.replace("---", "—").replace("--", "–")
        text = text.replace("``", "“").replace("''", "”")
        text = text.replace("~", "\u00a0")
        run = paragraph.add_run(text)
        set_run_font(run, "Cambria Math" if math else ("DejaVu Sans Mono" if code else "Times New Roman"), 10.5 if code else 13)
        run.bold = bold
        run.italic = italic or math
        if color:
            run.font.color.rgb = RGBColor.from_string(color)

    def _walk(self, paragraph, text: str, *, bold=False, italic=False, code=False, color=None):
        pos = 0
        plain = []

        def flush():
            if plain:
                self._plain(paragraph, "".join(plain), bold=bold, italic=italic, code=code, color=color)
                plain.clear()

        while pos < len(text):
            if text[pos] == "$":
                end = pos + 1
                while end < len(text) and text[end] != "$":
                    end += 1
                if end < len(text):
                    flush()
                    self._plain(paragraph, math_to_text(text[pos + 1:end]), bold=bold, italic=italic, code=False, color=color, math=True)
                    pos = end + 1
                    continue
            if text[pos] != "\\":
                plain.append(text[pos])
                pos += 1
                continue

            if pos + 1 < len(text) and text[pos + 1] in "%&_#$ {}":
                plain.append(text[pos + 1])
                pos += 2
                continue
            match = re.match(r"\\([A-Za-z]+)", text[pos:])
            if not match:
                plain.append("\\")
                pos += 1
                continue
            command = match.group(1)
            cend = pos + len(match.group(0))
            if command in ("noindent", "normalfont", "rmfamily", "bfseries", "small", "scriptsize", "selectfont"):
                pos = cend
                continue
            if command in ("statusmeasured", "statusimplemented", "statustarget", "statusmissing"):
                flush()
                status = {
                    "statusmeasured": ("Measured", "1B6E3C"),
                    "statusimplemented": ("Implemented", "0B4F6C"),
                    "statustarget": ("Target", "8A5A00"),
                    "statusmissing": ("Not measured", "9A1B1B"),
                }[command]
                self._plain(paragraph, status[0], bold=True, italic=False, code=False, color=status[1])
                pos = cend
                continue
            if command in ("cite", "ref", "textbf", "textit", "emph", "code", "path", "url", "text", "mathrm", "operatorname", "makecell"):
                arg_pos = cend
                if command == "makecell" and arg_pos < len(text) and text[arg_pos] == "[":
                    optional_end = text.find("]", arg_pos + 1)
                    if optional_end >= 0:
                        arg_pos = optional_end + 1
                arg, next_pos = brace_arg(text, arg_pos)
                if next_pos == cend:
                    plain.append(command)
                    pos = cend
                    continue
                flush()
                if command == "cite":
                    nums = [self.refs.citations.get(k.strip()) for k in arg.split(",")]
                    value = ", ".join(f"[{n}]" if n else f"[{k.strip()}]" for k, n in zip(arg.split(","), nums))
                    self._plain(paragraph, value, bold=bold, italic=italic, code=code, color=color)
                elif command == "ref":
                    self._plain(paragraph, self.refs.labels.get(arg, arg), bold=bold, italic=italic, code=code, color=color)
                else:
                    if command == "makecell":
                        arg = arg.replace(r"\\", " ")
                    self._walk(
                        paragraph,
                        arg,
                        bold=bold or command == "textbf",
                        italic=italic or command in ("textit", "emph"),
                        code=code or command in ("code", "path", "url"),
                        color=color,
                    )
                pos = next_pos
                continue
            if command in MATH_COMMANDS:
                plain.append(MATH_COMMANDS[command])
                pos = cend
                continue
            # Preserve unknown braced commands by retaining their human-visible argument.
            arg, next_pos = brace_arg(text, cend)
            if next_pos != cend:
                flush()
                self._walk(paragraph, arg, bold=bold, italic=italic, code=code, color=color)
                pos = next_pos
            else:
                pos = cend
        flush()


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_margins(cell, top=70, start=80, bottom=70, end=80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def split_top_level(text: str, delimiter: str) -> list[str]:
    parts, buf = [], []
    depth = 0
    pos = 0
    while pos < len(text):
        ch = text[pos]
        if ch == "\\" and pos + 1 < len(text) and text[pos + 1] in "&%_$#{}":
            buf.extend((ch, text[pos + 1]))
            pos += 2
            continue
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth = max(0, depth - 1)
        if depth == 0 and text.startswith(delimiter, pos):
            parts.append("".join(buf).strip())
            buf = []
            pos += len(delimiter)
            continue
        buf.append(ch)
        pos += 1
    parts.append("".join(buf).strip())
    return parts


def clean_table_markup(text: str) -> str:
    text = re.sub(r"\\caption\{.*?\}(?:\\label\{.*?\})?\\\\", "", text, flags=re.S)
    text = re.sub(r"\\label\{.*?\}", "", text)
    text = re.sub(r"\\(?:toprule|midrule|bottomrule|centering)\b", "", text)
    text = re.sub(r"\\(?:endfirsthead|endhead)\b", "", text)
    return text.strip()


def parse_rows(content: str, fr: bool = False) -> list[list[str]]:
    delimiter = "\\frtablerow" if fr else "\\\\"
    rows = []
    for raw in split_top_level(content, delimiter):
        raw = clean_table_markup(raw)
        if not raw:
            continue
        cols = split_top_level(raw, "&")
        if len(cols) > 1:
            rows.append(cols)
    return rows


def extract_table(block: str, env: str) -> tuple[str | None, str | None, list[list[str]]]:
    caption = label = None
    if env == "frtable":
        start = block.find("\\begin{frtable}") + len("\\begin{frtable}")
        caption, p1 = brace_arg(block, start)
        label, p2 = brace_arg(block, p1)
        end = block.rfind("\\end{frtable}")
        content = block[p2:end]
        return caption, label, [["Item", "Description"]] + parse_rows(content, fr=True)

    cap_match = re.search(r"\\caption\{", block)
    if cap_match:
        caption, _ = brace_arg(block, cap_match.end() - 1)
    label_match = re.search(r"\\label\{", block)
    if label_match:
        label, _ = brace_arg(block, label_match.end() - 1)

    inner_env = "tabularx" if "\\begin{tabularx}" in block else ("tabular" if "\\begin{tabular}" in block else "longtable")
    marker = f"\\begin{{{inner_env}}}"
    start = block.find(marker) + len(marker)
    if inner_env == "tabularx":
        _, start = brace_arg(block, start)  # width
        _, start = brace_arg(block, start)  # column spec
    else:
        _, start = brace_arg(block, start)  # column spec
    end = block.rfind(f"\\end{{{inner_env}}}")
    content = block[start:end]

    if "\\endfirsthead" in content and "\\endhead" in content:
        first, rest = content.split("\\endfirsthead", 1)
        _, body = rest.split("\\endhead", 1)
        first_rows = parse_rows(first)
        body_rows = parse_rows(body)
        rows = (first_rows[-1:] if first_rows else []) + body_rows
    else:
        rows = parse_rows(content)
    return caption, label, rows


def png_size(path: Path) -> tuple[int, int]:
    with path.open("rb") as handle:
        sig = handle.read(24)
    if sig[:8] != b"\x89PNG\r\n\x1a\n":
        return 1800, 1200
    return struct.unpack(">II", sig[16:24])


@dataclass
class Audit:
    headings: dict[str, int] = field(default_factory=lambda: {"chapter": 0, "section": 0, "subsection": 0, "subsubsection": 0})
    tables: int = 0
    captioned_tables: int = 0
    figures: int = 0
    equations: int = 0
    equation_fallbacks: int = 0
    bibliography: int = 0


class ReportExporter:
    def __init__(self, output: Path):
        self.output = output
        self.refs = build_reference_map()
        self.doc = Document()
        configure_styles(self.doc)
        add_update_fields_setting(self.doc)
        self.writer = InlineWriter(self.refs)
        self.audit = Audit()
        self.chapter = self.section = self.subsection = self.subsubsection = 0
        self.appendix = False
        self.appendix_index = 0
        self.table_seq = 0
        self.figure_seq = 0
        self.tmp = tempfile.TemporaryDirectory(prefix="perfin-docx-")
        self.tmpdir = Path(self.tmp.name)

    def add_front_title(self, title: str, listed: bool = True):
        p = self.doc.add_paragraph(style="Front Heading 1" if listed else "Unlisted Front Title")
        p.add_run(title.upper())
        return p

    def add_toc_pages(self):
        self.doc.add_page_break()
        self.add_front_title("Contents", listed=False)
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        add_field(p, ' TOC \\o "1-4" \\h \\z \\t "Front Heading 1,1" ', "Right-click and update this field to generate the table of contents.")

        self.doc.add_page_break()
        self.add_front_title("List of Tables")
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        add_field(p, ' TOC \\h \\z \\c "Table" ', "Right-click and update this field to generate the list of tables.")

        self.doc.add_page_break()
        self.add_front_title("List of Figures")
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        add_field(p, ' TOC \\h \\z \\c "Figure" ', "Right-click and update this field to generate the list of figures.")

    def add_caption(self, kind: str, caption: str):
        if kind == "Table":
            self.table_seq += 1
            sequence_number = self.table_seq
        else:
            self.figure_seq += 1
            sequence_number = self.figure_seq
        p = self.doc.add_paragraph(style="Caption")
        run = p.add_run(kind + " ")
        set_run_font(run, "Times New Roman", 13)
        run.bold = True
        add_field(p, f" SEQ {kind} \\* ARABIC ", str(sequence_number))
        run = p.add_run(". ")
        set_run_font(run, "Times New Roman", 13)
        self.writer.add(p, caption)
        return p

    def add_table(self, rows: list[list[str]], caption: str | None = None):
        if not rows:
            return
        # Guideline-report.md specifies No./Abbreviation/Full form.  Add the
        # required ordinal column while retaining every source abbreviation.
        header_plain = [re.sub(r"\\textbf\{([^}]*)\}", r"\1", x).strip() for x in rows[0]]
        if len(rows[0]) == 2 and header_plain == ["Abbreviation", "Meaning"]:
            rows = [["No.", "Abbreviation", "Full form"]] + [
                [str(index), row[0], row[1]]
                for index, row in enumerate(rows[1:], start=1)
            ]
        if caption:
            self.add_caption("Table", caption)
            self.audit.captioned_tables += 1
        width = max(len(row) for row in rows)
        table = self.doc.add_table(rows=len(rows), cols=width)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        for r_idx, values in enumerate(rows):
            row = table.rows[r_idx]
            if r_idx == 0:
                repeat_table_header(row)
            for c_idx in range(width):
                cell = row.cells[c_idx]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell)
                p = cell.paragraphs[0]
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = Pt(0)
                value = values[c_idx] if c_idx < len(values) else ""
                self.writer.add(p, value)
                if r_idx == 0:
                    shade_cell(cell, "D9EAF7")
                    for run in p.runs:
                        run.bold = True
        self.doc.add_paragraph().paragraph_format.space_after = Pt(0)
        self.audit.tables += 1

    def render_equation(self, tex: str, index: int) -> Path | None:
        tex_file = self.tmpdir / f"eq-{index:02d}.tex"
        pdf_file = self.tmpdir / f"eq-{index:02d}.pdf"
        png_file = self.tmpdir / f"eq-{index:02d}.png"
        source = (
            "\\documentclass[border=3pt]{standalone}\n"
            "\\usepackage{fontspec}\n"
            "\\usepackage{amsmath,amssymb}\n"
            "\\setmainfont{Times New Roman}\n"
            "\\begin{document}\n"
            "\\begin{minipage}{0.95\\textwidth}\n"
            "\\[\n" + tex.strip() + "\n\\]\n"
            "\\end{minipage}\n"
            "\\end{document}\n"
        )
        tex_file.write_text(source, encoding="utf-8")
        proc = subprocess.run(
            ["xelatex", "-interaction=nonstopmode", "-halt-on-error", tex_file.name],
            cwd=self.tmpdir,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            check=False,
        )
        if proc.returncode != 0 or not pdf_file.exists():
            return None
        proc = subprocess.run(
            ["mutool", "draw", "-q", "-r", "300", "-o", png_file.name, pdf_file.name, "1"],
            cwd=self.tmpdir,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            check=False,
        )
        return png_file if proc.returncode == 0 and png_file.exists() else None

    def add_equation(self, tex: str):
        self.audit.equations += 1
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        image = self.render_equation(tex, self.audit.equations)
        if image:
            px_w, px_h = png_size(image)
            width_in = min(6.0, px_w / 300.0)
            if px_h / 300.0 * (width_in / max(px_w / 300.0, 0.01)) > 2.8:
                width_in *= 2.8 / (px_h / 300.0 * width_in / max(px_w / 300.0, 0.01))
            p.add_run().add_picture(str(image), width=Inches(max(1.0, width_in)))
        else:
            self.audit.equation_fallbacks += 1
            run = p.add_run(math_to_text(tex))
            set_run_font(run, "Cambria Math", 13)
            run.italic = True

    def add_figure(self, filename: str, caption: str):
        # LaTeX deliberately places these wide diagrams on landscape pages.
        landscape = self.doc.add_section(WD_SECTION.NEW_PAGE)
        set_page_geometry(landscape, landscape=True)
        set_page_number_format(landscape, "decimal")
        add_page_number(landscape)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        png = ROOT / "figures/rendered" / (Path(filename).stem + ".png")
        if not png.exists():
            raise FileNotFoundError(png)
        px_w, px_h = png_size(png)
        max_w, max_h = 9.65, 5.75
        width = min(max_w, max_h * px_w / max(px_h, 1))
        p.add_run().add_picture(str(png), width=Inches(width))
        self.add_caption("Figure", caption)
        self.audit.figures += 1
        portrait = self.doc.add_section(WD_SECTION.NEW_PAGE)
        set_page_geometry(portrait, landscape=False)
        set_page_number_format(portrait, "decimal")
        add_page_number(portrait)

    def add_heading(self, level: str, title: str):
        if level == "chapter":
            self.section = self.subsection = self.subsubsection = 0
            if self.appendix:
                self.appendix_index += 1
                number = chr(64 + self.appendix_index)
                text = f"APPENDIX {number}: {title.upper()}"
            else:
                self.chapter += 1
                text = f"CHAPTER {self.chapter}: {title.upper()}"
            style = "Heading 1"
        elif level == "section":
            self.section += 1
            self.subsection = self.subsubsection = 0
            prefix = chr(64 + self.appendix_index) if self.appendix else str(self.chapter)
            text = f"{prefix}.{self.section}. {title.upper()}"
            style = "Heading 2"
        elif level == "subsection":
            self.subsection += 1
            self.subsubsection = 0
            prefix = chr(64 + self.appendix_index) if self.appendix else str(self.chapter)
            text = f"{prefix}.{self.section}.{self.subsection}. {title}"
            style = "Heading 3"
        else:
            self.subsubsection += 1
            prefix = chr(64 + self.appendix_index) if self.appendix else str(self.chapter)
            text = f"{prefix}.{self.section}.{self.subsection}.{self.subsubsection}. {title}"
            style = "Heading 4"
        text = text.replace("---", "—").replace("--", "–")
        p = self.doc.add_paragraph(style=style)
        p.add_run(text)
        self.audit.headings[level] += 1

    def add_paragraph(self, text: str, style: str | None = None):
        no_indent = text.lstrip().startswith("\\noindent")
        text = re.sub(r"\\vspace\{[^}]+\}", "", text)
        text = text.replace("\\par", " ")
        text = re.sub(r"\s+", " ", text).strip()
        if not text:
            return
        p = self.doc.add_paragraph(style=style)
        if no_indent or style in ("List Bullet", "List Number", "Code Block"):
            p.paragraph_format.first_line_indent = Cm(0)
        self.writer.add(p, text)

    def add_list(self, block: str, ordered: bool):
        env = "enumerate" if ordered else "itemize"
        body = block.split(f"\\begin{{{env}}}", 1)[1].rsplit(f"\\end{{{env}}}", 1)[0]
        parts = re.split(r"(?m)^\s*\\item\s*", body)
        for item in parts[1:]:
            self.add_paragraph(" ".join(item.splitlines()), "List Number" if ordered else "List Bullet")

    def add_verbatim(self, block: str):
        body = block.split("\\begin{verbatim}", 1)[1].rsplit("\\end{verbatim}", 1)[0].strip("\n")
        p = self.doc.add_paragraph(style="Code Block")
        for idx, line in enumerate(body.splitlines()):
            run = p.add_run(("\n" if idx else "") + line)
            set_run_font(run, "DejaVu Sans Mono", 10.5)

    def add_bibliography(self, block: str):
        self.doc.add_page_break()
        self.add_front_title("References")
        body = block.split("\\begin{thebibliography}", 1)[1]
        body = body[body.find("}") + 1 :].rsplit("\\end{thebibliography}", 1)[0]
        items = list(re.finditer(r"\\bibitem\{([^}]+)\}", body))
        for idx, match in enumerate(items):
            end = items[idx + 1].start() if idx + 1 < len(items) else len(body)
            entry = " ".join(body[match.end():end].splitlines()).strip()
            p = self.doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(-0.8)
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_after = Pt(5)
            n = self.refs.citations.get(match.group(1), idx + 1)
            run = p.add_run(f"[{n}] ")
            set_run_font(run, "Times New Roman", 13)
            self.writer.add(p, entry)
            self.audit.bibliography += 1

    def process_file(self, path: Path, *, suppress_first_front_break: bool = False):
        text = strip_comments(path.read_text(encoding="utf-8"))
        lines = text.splitlines()
        paragraph = []
        first_front = True

        def flush():
            if paragraph:
                self.add_paragraph(" ".join(paragraph))
                paragraph.clear()

        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                flush()
                i += 1
                continue
            heading = re.match(r"\\(chapter|section|subsection|subsubsection)\{", line)
            if heading:
                flush()
                title, _ = brace_arg(line, heading.end() - 1)
                self.add_heading(heading.group(1), title)
                i += 1
                continue
            front = re.match(r"\\frontchapter\{", line)
            if front:
                flush()
                title, _ = brace_arg(line, front.end() - 1)
                if not (suppress_first_front_break and first_front):
                    self.doc.add_page_break()
                self.add_front_title(title)
                first_front = False
                i += 1
                continue
            figure = re.match(r"\\widereportfigure\{", line)
            if figure:
                flush()
                filename, p1 = brace_arg(line, figure.end() - 1)
                caption, p2 = brace_arg(line, p1)
                _, _ = brace_arg(line, p2)
                self.add_figure(filename, caption)
                i += 1
                continue
            begin = re.match(r"\\begin\{([^}]+)\}", line)
            if begin:
                flush()
                env = begin.group(1)
                block_lines = [lines[i]]
                i += 1
                while i < len(lines):
                    block_lines.append(lines[i])
                    if f"\\end{{{env}}}" in lines[i]:
                        i += 1
                        break
                    i += 1
                block = "\n".join(block_lines)
                if env in ("table", "longtable", "frtable"):
                    caption, _, rows = extract_table(block, env)
                    self.add_table(rows, caption)
                elif env == "equation":
                    body = block.split("\\begin{equation}", 1)[1].rsplit("\\end{equation}", 1)[0]
                    self.add_equation(body)
                elif env in ("enumerate", "itemize"):
                    self.add_list(block, env == "enumerate")
                elif env == "verbatim":
                    self.add_verbatim(block)
                elif env == "thebibliography":
                    self.add_bibliography(block)
                continue
            if line.startswith("\\appendix"):
                flush()
                self.appendix = True
                appendix_section = self.doc.add_section(WD_SECTION.NEW_PAGE)
                set_page_geometry(appendix_section)
                set_page_number_format(appendix_section, "decimal", 1)
                add_page_number(appendix_section)
                i += 1
                continue
            if line.startswith(("\\label", "\\cleardoublepage", "\\phantomsection", "\\addcontentsline", "\\setcounter", "\\startappendixpages")):
                flush()
                i += 1
                continue
            paragraph.append(line)
            i += 1
        flush()

    def export(self):
        set_page_geometry(self.doc.sections[0])
        add_cover(self.doc)
        preliminary = self.doc.add_section(WD_SECTION.NEW_PAGE)
        set_page_geometry(preliminary)
        set_page_number_format(preliminary, "lowerRoman", 1)
        add_page_number(preliminary)

        self.process_file(ROOT / "frontmatter/en/abstract.tex", suppress_first_front_break=True)
        self.add_toc_pages()
        self.process_file(ROOT / "frontmatter/en/abbreviations.tex")

        main = self.doc.add_section(WD_SECTION.NEW_PAGE)
        set_page_geometry(main)
        set_page_number_format(main, "decimal", 1)
        add_page_number(main)
        for path in EN_FILES[2:]:
            self.process_file(path)

        self.doc.core_properties.title = META["projecttitle"]
        self.doc.core_properties.author = META["studentname"]
        self.doc.core_properties.subject = "Fundamental-topics project report — English DOCX edition"
        self.doc.core_properties.keywords = "personal finance, LLM, PostgreSQL, analytics, PERFIN"
        self.output.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(self.output)
        self.tmp.cleanup()
        return self.audit


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "output",
        nargs="?",
        type=Path,
        default=ROOT / "PERFIN_Report_EN.docx",
        help="output DOCX path",
    )
    args = parser.parse_args()
    exporter = ReportExporter(args.output.resolve())
    audit = exporter.export()
    print(f"Wrote: {args.output.resolve()}")
    print(f"Headings: {audit.headings}")
    print(f"Tables: {audit.tables} total, {audit.captioned_tables} captioned")
    print(f"Figures: {audit.figures}")
    print(f"Equations: {audit.equations} ({audit.equation_fallbacks} text fallbacks)")
    print(f"References: {audit.bibliography}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
